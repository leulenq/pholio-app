#!/usr/bin/env python3
"""Generate the Pholio-styled pre-production founder/status DOCX.

The document intentionally mirrors the public legal-page rhythm used on
pholio-landing: legal/compliance label, large serif title, last updated / effective
metadata, table of contents, numbered legal-document sections, warm paper palette,
gold rules, quiet callout boxes, and a footer disclaimer.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs/business/pre-production-founder-status-report.md"
OUTPUT = ROOT / "docs/business/pre-production-founder-status-report.docx"

INK = "1A1815"
MUTED = "6B6560"
GOLD = "B8956A"
GOLD_DARK = "8E714D"
PAPER = "FAF8F5"
SURFACE = "FFFFFF"
LINE = "E8DED2"
CALLOUT = "F4EDE3"
DANGER = "7A2E2E"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_bottom_border(paragraph, color: str = GOLD, size: str = "14") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_page_background(section, fill: str = PAPER) -> None:
    sect_pr = section._sectPr
    bg = sect_pr.find(qn("w:background"))
    if bg is None:
        bg = OxmlElement("w:background")
        sect_pr.insert(0, bg)
    bg.set(qn("w:color"), fill)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    set_page_background(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Inter"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Inter")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [
        ("Heading 1", 19, 18, 7),
        ("Heading 2", 13, 14, 5),
        ("Heading 3", 10.5, 9, 3),
    ]:
        style = styles[name]
        style.font.name = "Playfair Display" if name == "Heading 1" else "Inter"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = name != "Heading 1"
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_run_with_markup(paragraph, text: str, color: str = INK) -> None:
    # Handles **bold** fragments from the Markdown source.
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        run = paragraph.add_run(clean)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.name = "Inter"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Inter")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GOLD_DARK)
    run.font.spacing = Pt(1.4)


def add_title_page(doc: Document) -> None:
    add_label(doc, "Legal & Compliance")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(7)
    run = title.add_run("Pholio Pre-Production Founder Status Report")
    run.font.name = "Playfair Display"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Playfair Display")
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    srun = subtitle.add_run("Cofounder meeting document for company, launch, legal, infrastructure, and Studio+ decisions before production.")
    srun.font.size = Pt(10.5)
    srun.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    rows = [
        ("Last Updated", "July 6, 2026"),
        ("Effective Planning Date", "July 6, 2026"),
        ("Audience", "Pholio cofounders, business/legal advisors, and Fashion Week Brooklyn launch partner"),
        ("Document Type", "Business/status report — not a product requirements spec"),
    ]
    for row, (key, value) in zip(meta.rows, rows):
        for cell in row.cells:
            set_cell_shading(cell, SURFACE)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].width = Inches(1.75)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(key)
        r0.bold = True
        r0.font.size = Pt(8.5)
        r0.font.color.rgb = RGBColor.from_string(GOLD_DARK)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        r1.font.size = Pt(8.8)
        r1.font.color.rgb = RGBColor.from_string(INK)

    rule = doc.add_paragraph()
    add_bottom_border(rule)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Inches(0.3)
    summary.paragraph_format.right_indent = Inches(0.3)
    add_run_with_markup(
        summary,
        "Recommended posture: launch as a controlled partner pilot, not an unrestricted public production launch, until counsel, infrastructure, support, and paid-service readiness are cleared.",
        MUTED,
    )


def collect_toc(markdown: str) -> list[str]:
    items = []
    for line in markdown.splitlines():
        if line.startswith("## ") and not line.startswith("## 14. Source"):
            text = line[3:].strip()
            if text not in {"1. Executive summary"}:
                items.append(text)
    return items


def add_toc(doc: Document, markdown: str) -> None:
    h = doc.add_heading("Table of Contents", level=1)
    add_bottom_border(h, LINE, "8")
    for idx, item in enumerate(collect_toc(markdown), start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(item)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_callout(doc: Document, lines: list[str], warning: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8F2EA" if not warning else "F6EAEA")
    set_cell_border(cell, GOLD if not warning else DANGER, "10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run_with_markup(p, " ".join(lines).replace("> ", ""), DANGER if warning else INK)


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return
    table = doc.add_table(rows=len(parsed), cols=max(len(r) for r in parsed))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, cells in enumerate(parsed):
        for c_idx, value in enumerate(cells):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, CALLOUT if r_idx == 0 else SURFACE)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run_with_markup(p, value)
            for run in p.runs:
                run.font.size = Pt(8.2)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(GOLD_DARK)


def add_source_note(doc: Document) -> None:
    h = doc.add_heading("Questions About This Report?", level=2)
    p = doc.add_paragraph()
    add_run_with_markup(
        p,
        "This document is an operator's planning brief, not legal advice. Confirm pricing before purchase and route founder ownership, immigration, privacy, subscription, and partner-launch questions to qualified counsel.",
        MUTED,
    )
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run("privacy@pholio.studio")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GOLD_DARK)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("© 2026 Pholio Studio, Inc. • Built for talent. Trusted by agencies. • This document does not constitute legal advice.")
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def render_body(doc: Document, markdown: str) -> None:
    skip_prefixes = {
        "# Pholio Pre-Production Founder Status Report",
        "**Audience:**",
        "**Prepared:**",
        "**Document type:**",
        "**Decision window:**",
        "----",
    }
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line or line == "---" or any(line.startswith(prefix) for prefix in skip_prefixes):
            idx += 1
            continue

        if line.startswith("|"):
            rows = []
            while idx < len(lines) and lines[idx].startswith("|"):
                rows.append(lines[idx])
                idx += 1
            add_markdown_table(doc, rows)
            continue

        if line.startswith(">"):
            block = []
            warning = "not legal advice" in line.lower() or "only if" in line.lower()
            while idx < len(lines) and lines[idx].startswith(">"):
                block.append(lines[idx])
                idx += 1
            add_callout(doc, block, warning=warning)
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            paragraph = doc.add_heading(text, level=1)
            add_bottom_border(paragraph, LINE, "8")
            idx += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            idx += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            idx += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_run_with_markup(p, re.sub(r"^\d+\.\s+", "", line))
            idx += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.28)
            add_run_with_markup(p, line[2:])
            idx += 1
            continue

        p = doc.add_paragraph()
        add_run_with_markup(p, line)
        idx += 1


def main() -> None:
    markdown = SOURCE.read_text()
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    doc.add_page_break()
    add_toc(doc, markdown)
    doc.add_page_break()
    render_body(doc, markdown)
    add_source_note(doc)
    add_footer(doc)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
